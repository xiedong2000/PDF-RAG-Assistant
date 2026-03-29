import io

import chromadb
import pdfplumber
import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv()

st.set_page_config(page_title="Document RAG Assistant", page_icon="📄")

st.title("📄 Document RAG Assistant")

client = OpenAI()

ALLOWED_TYPES = ["pdf", "docx", "txt"]

uploaded_files = st.file_uploader(
    "Upload one or more PDF, DOCX, or TXT files",
    type=ALLOWED_TYPES,
    accept_multiple_files=True,
)

if uploaded_files:

    def extract_document_text(uploaded_file) -> str:
        name = (uploaded_file.name or "").lower()
        uploaded_file.seek(0)

        if name.endswith(".pdf"):
            text = ""
            with pdfplumber.open(uploaded_file) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text
                    except Exception as e:
                        st.warning(
                            f"{uploaded_file.name} — could not extract text from page "
                            f"{page_num + 1}: {str(e)}"
                        )
            return text

        if name.endswith(".docx"):
            try:
                from docx import Document
            except ImportError:
                st.error(
                    "DOCX support needs the **python-docx** package. "
                    "Run: `pip install python-docx` or `pip install -r requirements.txt`."
                )
                return ""
            data = uploaded_file.read()
            doc = Document(io.BytesIO(data))
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            return "\n".join(parts)

        if name.endswith(".txt"):
            raw = uploaded_file.read()
            return raw.decode("utf-8", errors="replace")

        st.warning(f'Unsupported file type: "{uploaded_file.name}"')
        return ""

    chunks = []
    chunk_ids = []
    chunk_sources = []
    files_indexed = 0
    for uploaded_file in uploaded_files:
        text = extract_document_text(uploaded_file)
        if not text.strip():
            st.warning(f'Skipping "{uploaded_file.name}": no extractable text.')
            continue
        labeled = f"[Source: {uploaded_file.name}]\n{text}"
        file_chunks = [labeled[i : i + 500] for i in range(0, len(labeled), 500)]
        for j, chunk in enumerate(file_chunks):
            chunks.append(chunk)
            chunk_ids.append(f"f{files_indexed}_c{j}")
            chunk_sources.append(uploaded_file.name)
        files_indexed += 1

    if not chunks:
        st.error(
            "Could not extract any text from the uploaded files. Try different documents."
        )
        st.stop()

    st.write(
        f"Indexed {files_indexed} file(s) ({len(chunks)} chunk(s) total)."
    )

    chroma_client = chromadb.Client()

    try:
        chroma_client.delete_collection(name="user_documents")
    except Exception:
        pass

    collection = chroma_client.get_or_create_collection(name="user_documents")

    for chunk, cid, source in zip(chunks, chunk_ids, chunk_sources):
        embedding = client.embeddings.create(
            model="text-embedding-3-small",
            input=chunk
        ).data[0].embedding

        collection.add(
            documents=[chunk],
            embeddings=[embedding],
            ids=[cid],
            metadatas=[{"source": source}]
        )

    question = st.text_input("Ask a question about your document(s)")

    if question:

        q_embedding = client.embeddings.create(
            model="text-embedding-3-small",
            input=question
        ).data[0].embedding

        results = collection.query(
            query_embeddings=[q_embedding],
            n_results=min(len(chunks), 5),
            include=["documents", "metadatas", "distances"]
        )

        retrieved_docs = results["documents"][0]
        retrieved_meta = results["metadatas"][0]
        retrieved_distances = results["distances"][0]

        st.subheader("Retrieved chunks")
        for idx, (doc, meta, distance) in enumerate(zip(retrieved_docs, retrieved_meta, retrieved_distances), start=1):
            source_name = meta.get("source", "unknown")
            with st.expander(f"{idx}. {source_name} — distance {distance:.4f}"):
                st.write(doc)

        context = "\n".join(retrieved_docs)

        prompt = f"""
        Use only the provided context to answer the question.
        If the answer is not contained in the context, say you do not know.

        Context:
        {context}

        Question:
        {question}
        """

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}]
        )

        st.subheader("Answer")
        st.write(response.choices[0].message.content)
